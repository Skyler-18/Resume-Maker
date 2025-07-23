def add_hyperlink(paragraph, url, text, bold=False, italic=False):
    # Adds a clickable hyperlink to a paragraph (python-docx workaround)
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Style: black and not underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'none')
    rPr.append(u)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hr(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    p._p.get_or_add_pPr().append(OxmlElement('w:pBdr'))
    p.paragraph_format.space_after = Pt(12)  # 1 line after HR for section spacing

def generate_docx(form_data, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    # Set very low top/bottom and reduced left/right margins
    for section in doc.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)  # Reduce space after heading

    def add_bullets(items):
        for item in items:
            para = doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.left_indent = Inches(0.25)

    # Name
    name_p = doc.add_paragraph(form_data.get('your_name', '').upper())
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.runs[0].font.size = Pt(18)
    name_p.runs[0].bold = True

    # Contact (with hyperlinks)
    c_p = doc.add_paragraph()
    c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = True
    def add_sep():
        nonlocal first
        if not first:
            c_p.add_run(' | ')
        first = False
    if form_data.get('email'):
        add_sep()
        email = form_data['email']
        add_hyperlink(c_p, f"mailto:{email}", f"Email: {email}")
    if form_data.get('linkedin'):
        add_sep()
        add_hyperlink(c_p, form_data['linkedin'], "LinkedIn")
    if form_data.get('github'):
        add_sep()
        add_hyperlink(c_p, form_data['github'], "GitHub")
    if form_data.get('website'):
        add_sep()
        add_hyperlink(c_p, form_data['website'], "Website")
    if form_data.get('phone'):
        add_sep()
        c_p.add_run(f"Phone: {form_data['phone']}")
    c_p.paragraph_format.space_after = Pt(12)

    # Summary
    if form_data.get('summary'):
        add_heading('Summary')
        doc.add_paragraph(form_data['summary']).paragraph_format.space_after = Pt(12)
        add_hr(doc)

    # Education
    from utils.pdf_generator import parse_education, parse_experience, parse_projects, parse_skills, parse_certifications, parse_activities, parse_publications, parse_custom_section
    education = parse_education(form_data)
    if education:
        add_heading('Education')
        table = doc.add_table(rows=0, cols=2)
        table.autofit = True
        for edu in education:
            row = table.add_row().cells
            row[0].text = edu['institute']
            row[0].paragraphs[0].runs[0].bold = True
            # City, State (normal)
            p_city = row[1].paragraphs[0]
            p_city.add_run(f"{edu['city']}, {edu['state']}")
            row2 = table.add_row().cells
            # Degree and branch in italics
            p_deg = row2[0].paragraphs[0]
            run_deg = p_deg.add_run(f"{edu['degree']} — {edu['branch']}")
            run_deg.italic = True
            # Date duration in italics
            p_date = row2[1].paragraphs[0]
            run_date = p_date.add_run(f"{edu['start']} – {edu['end']}")
            run_date.italic = True
            # Do NOT clear row2[1].text here
        # Remove extra paragraph, just add HR for section separation
        add_hr(doc)
        # Add 1 line spacing after section
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Experience
    experience = parse_experience(form_data)
    if experience:
        add_heading('Experience')
        for exp in experience:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            row = table.add_row().cells
            # Experience Role with hyperlink if present
            p = row[0].paragraphs[0]
            if exp.get('role_link'):
                add_hyperlink(p, exp['role_link'], exp['role'], bold=True)
            else:
                run = p.add_run(exp['role'])
                run.bold = True
            # Date duration (normal)
            p_date = row[1].paragraphs[0]
            p_date.add_run(f"{exp['start']} – {exp['end']}")
            # Do NOT clear row[1].text here
            row2 = table.add_row().cells
            # Organization in italics
            p_org = row2[0].paragraphs[0]
            run_org = p_org.add_run(f"{exp['org']}")
            run_org.italic = True
            # City, State in italics
            p_city = row2[1].paragraphs[0]
            run_city = p_city.add_run(f"{exp['city']}, {exp['state']}")
            run_city.italic = True
            # Do NOT clear row2[1].text here
            if exp['bullets']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                add_bullets(exp['bullets'])
        add_hr(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Projects
    projects = parse_projects(form_data)
    if projects:
        add_heading('Projects')
        for proj in projects:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            row = table.add_row().cells
            # Project name with hyperlink if present
            p = row[0].paragraphs[0]
            if proj.get('link'):
                add_hyperlink(p, proj['link'], proj['name'], bold=True)
            else:
                run = p.add_run(proj['name'])
                run.bold = True
            # Date duration (normal)
            p_date = row[1].paragraphs[0]
            p_date.add_run(f"{proj['start']} – {proj['end']}")
            # Do NOT clear row[1].text here
            if proj['tech']:
                tech_p = doc.add_paragraph()
                tech_p.paragraph_format.left_indent = Inches(0.25)
                run_tech = tech_p.add_run(f"{proj['tech']}")
                run_tech.italic = True
            if proj['bullets']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                add_bullets(proj['bullets'])
        add_hr(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Technical Skills
    skills = parse_skills(form_data)
    if skills:
        add_heading('Technical Skills')
        for skill in skills:
            p = doc.add_paragraph()
            run_heading = p.add_run(f"{skill['heading']}: ")
            run_heading.bold = True
            p.add_run(f"{skill['skills']}")
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        add_hr(doc)

    # Certifications
    certifications = parse_certifications(form_data)
    if certifications:
        add_heading('Certifications')
        for cert in certifications:
            p = doc.add_paragraph()
            # Certification name with hyperlink if present
            if cert.get('link'):
                add_hyperlink(p, cert['link'], cert['name'], bold=True)
            else:
                run = p.add_run(cert['name'])
                run.bold = True
            run_org = p.add_run(f" — {cert['org']}")
            run_org.italic = True
        add_hr(doc)

    # Co-curricular Activities
    co_curricular = parse_activities(form_data, 'co')
    if co_curricular:
        add_heading('Co-curricular Activities')
        for act in co_curricular:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            row = table.add_row().cells
            # Activity name with hyperlink if present
            p = row[0].paragraphs[0]
            if act.get('link'):
                add_hyperlink(p, act['link'], act['activity'], bold=True)
            else:
                run = p.add_run(act['activity'])
                run.bold = True
            # Date duration (normal)
            p_date = row[1].paragraphs[0]
            p_date.add_run(f"{act['start']} – {act['end']}")
            # Do NOT clear row[1].text here
            row2 = table.add_row().cells
            # Organization in italics
            p_org = row2[0].paragraphs[0]
            run_org = p_org.add_run(f"{act['org']}")
            run_org.italic = True
            # City, State in italics
            p_city = row2[1].paragraphs[0]
            run_city = p_city.add_run(f"{act['city']}, {act['state']}")
            run_city.italic = True
            # Do NOT clear row2[1].text here
            if act['bullets']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                add_bullets(act['bullets'])
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        add_hr(doc)

    # Extra-curricular Activities
    extra_curricular = parse_activities(form_data, 'extra')
    if extra_curricular:
        add_heading('Extra-curricular Activities')
        for act in extra_curricular:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            row = table.add_row().cells
            p = row[0].paragraphs[0]
            if act.get('link'):
                add_hyperlink(p, act['link'], act['activity'], bold=True)
            else:
                run = p.add_run(act['activity'])
                run.bold = True
            # Date duration in italics
            p_date = row[1].paragraphs[0]
            run_date = p_date.add_run(f"{act['start']} – {act['end']}")
            # Do NOT clear row[1].text here
            row2 = table.add_row().cells
            # Organization in italics
            p_org = row2[0].paragraphs[0]
            run_org = p_org.add_run(f"{act['org']}")
            run_org.italic = True
            # City, State in italics
            p_city = row2[1].paragraphs[0]
            run_city = p_city.add_run(f"{act['city']}, {act['state']}")
            run_city.italic = True
            # Do NOT clear row2[1].text here
            if act['bullets']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                add_bullets(act['bullets'])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        add_hr(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Publications
    publications = parse_publications(form_data)
    if publications:
        add_heading('Publications')
        for pub in publications:
            p = doc.add_paragraph()
            # Publication title with hyperlink if present
            if pub.get('link'):
                add_hyperlink(p, pub['link'], f'"{pub["title"]}"', bold=True)
            else:
                run = p.add_run(f'"{pub["title"]}"')
                run.bold = True
            p.add_run(f' {pub["author"]}, {pub["journal"]}, {pub["date"]}')
        add_hr(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Custom Section
    custom_section = parse_custom_section(form_data)
    if custom_section:
        add_heading(custom_section['heading'])
        add_bullets(custom_section['points'])

    doc.save(output_path) 